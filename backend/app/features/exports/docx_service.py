from io import BytesIO
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


TEAL = RGBColor(0, 245, 212)
MAGENTA = RGBColor(255, 0, 127)
NAVY = RGBColor(8, 16, 36)
DARK_TEXT = RGBColor(22, 26, 43)
MUTED = RGBColor(102, 112, 133)


def _safe(value, fallback="—"):
    if value is None:
        return fallback

    value = str(value).strip()

    return value or fallback


def _format_created_at(value):
    if not value:
        return ""

    try:
        parsed = datetime.fromisoformat(
            value.replace("Z", "+00:00")
        )

        return parsed.strftime(
            "%b %d, %Y • %I:%M %p"
        )

    except (ValueError, AttributeError):
        return str(value)


def _result_sections(tool_id, result):
    if tool_id == "clarity":
        return [
            (
                "Clarity Score",
                f"{result.get('score', '—')}/100"
            ),
            (
                "Status",
                result.get("label")
            ),
            (
                "Recommendation",
                result.get("recommendation")
            ),
            (
                "Refined Executive Copy",
                result.get("refined_text")
            ),
        ]

    if tool_id == "kpi-cleaner":
        return [
            (
                "Issues Found",
                result.get("issues_found")
            ),
            (
                "Status",
                result.get("label")
            ),
            (
                "Cleaned KPI Output",
                result.get("result")
            ),
        ]

    if tool_id == "insights":
        return [
            (
                "Primary Insight",
                result.get("primary_insight")
            ),
            (
                "Executive Implication",
                result.get("so_what")
            ),
            (
                "Recommended Action",
                result.get("recommended_action")
            ),
            (
                "Suggested Executive Title",
                result.get("executive_title")
            ),
            (
                "Suggested Visualization",
                result.get("chart_suggestion")
            ),
        ]

    if tool_id == "dashboard":
        return [
            (
                "Executive Summary",
                result.get("executive_summary")
            ),
            (
                "Performance Drivers",
                result.get("performance_drivers")
            ),
            (
                "Risks & Watch Items",
                result.get("risks")
            ),
            (
                "Recommended Action",
                result.get("recommended_action")
            ),
            (
                "Outlook",
                result.get("outlook")
            ),
        ]

    if tool_id == "executive-memo":
        return [
            (
                "Executive Summary",
                result.get("summary")
            ),
            (
                "Background",
                result.get("background")
            ),
            (
                "Key Findings",
                result.get("findings")
            ),
            (
                "Business Impact",
                result.get("impact")
            ),
            (
                "Recommendations",
                result.get("recommendations")
            ),
            (
                "Next Steps",
                result.get("next_steps")
            ),
        ]

    if tool_id == "kpi-health":
        return [
            (
                "Overall Health Score",
                f"{result.get('overall_score', '—')}/100"
            ),
            (
                "Executive Assessment",
                result.get("summary")
            ),
            (
                "Strengths",
                result.get("strengths", [])
            ),
            (
                "Concerns",
                result.get("concerns", [])
            ),
            (
                "Recommendations",
                result.get("recommendations", [])
            ),
        ]

    return [
        (
            "Analysis Result",
            str(result)
        )
    ]


def _set_cell_border(cell, color="D9E0EB"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    borders = tcPr.first_child_found_in(
        "w:tcBorders"
    )

    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)

    for edge in [
        "top",
        "left",
        "bottom",
        "right"
    ]:
        tag = "w:" + edge

        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), color)


def _add_section_card(document, heading, value):
    table = document.add_table(
        rows=1,
        cols=1
    )

    table.autofit = True

    cell = table.cell(0, 0)
    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )

    _set_cell_border(cell)

    heading_paragraph = cell.paragraphs[0]

    heading_run = heading_paragraph.add_run(
        _safe(heading).upper()
    )

    heading_run.bold = True
    heading_run.font.size = Pt(8)
    heading_run.font.color.rgb = TEAL

    if isinstance(value, list):
        for item in value:
            paragraph = cell.add_paragraph(
                style=None
            )

            paragraph.style = (
                document.styles["Normal"]
            )

            run = paragraph.add_run(
                f"• {_safe(item)}"
            )

            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_TEXT

    else:
        paragraph = cell.add_paragraph()

        run = paragraph.add_run(
            _safe(value)
        )

        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK_TEXT

    document.add_paragraph()


def _add_footer(section):
    footer = section.footer

    paragraph = footer.paragraphs[0]

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = paragraph.add_run(
        "Moon Onyx Labs • Executive Insight Engine"
    )

    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def build_analysis_docx(payload):
    document = Document()

    section = document.sections[0]

    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    _add_footer(section)

    normal = document.styles["Normal"]

    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = DARK_TEXT

    # -------------------------
    # BRAND
    # -------------------------

    brand = document.add_paragraph()

    brand_run = brand.add_run(
        "MOON ONYX LABS"
    )

    brand_run.bold = True
    brand_run.font.size = Pt(9)
    brand_run.font.color.rgb = TEAL

    report_type = document.add_paragraph()

    report_run = report_type.add_run(
        "EXECUTIVE INTELLIGENCE REPORT"
    )

    report_run.bold = True
    report_run.font.size = Pt(8)
    report_run.font.color.rgb = MAGENTA

    # -------------------------
    # TITLE
    # -------------------------

    title = document.add_paragraph()

    title_run = title.add_run(
        _safe(payload.title)
    )

    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = NAVY

    # -------------------------
    # METADATA
    # -------------------------

    meta = [
        _safe(payload.toolName, ""),
        _safe(payload.status, ""),
        _format_created_at(payload.createdAt),
    ]

    meta = [
        item for item in meta if item
    ]

    if meta:
        meta_paragraph = document.add_paragraph()

        meta_run = meta_paragraph.add_run(
            "  •  ".join(meta)
        )

        meta_run.font.size = Pt(8.5)
        meta_run.font.color.rgb = MUTED

    document.add_paragraph()

    # -------------------------
    # ANALYSIS
    # -------------------------

    analysis_label = document.add_paragraph()

    label_run = analysis_label.add_run(
        "EXECUTIVE ANALYSIS"
    )

    label_run.bold = True
    label_run.font.size = Pt(8)
    label_run.font.color.rgb = TEAL

    analysis_heading = document.add_paragraph()

    analysis_heading_run = (
        analysis_heading.add_run(
            "Analysis Result"
        )
    )

    analysis_heading_run.bold = True
    analysis_heading_run.font.size = Pt(14)
    analysis_heading_run.font.color.rgb = NAVY

    for heading, value in _result_sections(
        payload.toolId,
        payload.result
    ):
        if value in (None, "", []):
            continue

        _add_section_card(
            document,
            heading,
            value
        )

    # -------------------------
    # SOURCE MATERIAL
    # -------------------------

    if payload.input:
        document.add_paragraph()

        source_label = document.add_paragraph()

        source_label_run = source_label.add_run(
            "SOURCE MATERIAL"
        )

        source_label_run.bold = True
        source_label_run.font.size = Pt(8)
        source_label_run.font.color.rgb = TEAL

        source_heading = document.add_paragraph()

        source_heading_run = (
            source_heading.add_run(
                "Original Input"
            )
        )

        source_heading_run.bold = True
        source_heading_run.font.size = Pt(14)
        source_heading_run.font.color.rgb = NAVY

        source_table = document.add_table(
            rows=1,
            cols=1
        )

        source_cell = source_table.cell(
            0,
            0
        )

        _set_cell_border(
            source_cell,
            color="00F5D4"
        )

        source_paragraph = (
            source_cell.paragraphs[0]
        )

        source_run = source_paragraph.add_run(
            _safe(payload.input)
        )

        source_run.font.size = Pt(9.5)
        source_run.font.color.rgb = DARK_TEXT

    document.add_paragraph()

    prepared = document.add_paragraph()

    prepared_run = prepared.add_run(
        "Prepared by Moon Onyx Labs for executive "
        "review and decision support."
    )

    prepared_run.italic = True
    prepared_run.font.size = Pt(8)
    prepared_run.font.color.rgb = MUTED

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer